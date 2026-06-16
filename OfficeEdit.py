import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from typing import List, Dict, Any
import os
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement


"""
一、段落类型（type="paragraph"）
{
    "type": "paragraph",  # 固定值：段落类型
    "config": {           # 段落级全局配置（作用于整个段落）
        "line_spacing": 1,        # int/float 行间距，默认1（1=单倍、1.5=1.5倍、2=双倍、数字=自定义磅值）
        "first_line_indent": 0,   # int 首行缩进（单位：字符，1=1字符、2=2字符，默认0）
        "space_before": 0,        # float 段前间距（单位：厘米，默认0）
        "space_after": 0,         # float 段后间距（单位：厘米，默认0）
        "type": "paragraph"       # 冗余字段，用于快速识别类型（可省略）
    },
    "content": [  # 段落内文本片段列表（支持多段不同样式的文本拼接）
        {
            "text": "这是一段写入的内容",     # str 文本内容（必填，空字符串不显示）
            "color": "#000000",             # str 字体颜色（16进制，支持#FF0000/FF0000/#f00/f00，默认#000000）
            "name": "宋体",                  # str 字体名称（系统已安装字体，默认宋体）
            "size": 12,                     # int 字号（单位：磅，默认12）
            "bold": "False",                # str 加粗（"True"/"False"，默认"False"）
            "italic": "False",              # str 斜体（"True"/"False"，默认"False"）
            "underline": "False",           # str 下划线（"False"/"single"/"double"/"dotted"/"dash"，默认"False"）
            "alignment": "left"            # str 对齐方式（"left"/"center"/"right"/"justify"，默认"left"）
        }
    ]
}

二、表格类型（type="table"）
{
    "type": "table",  # 固定值：表格类型
    "config": {
        "rows": 2,                # int 行数（必填，默认1）
        "cols": 3,                # int 列数（必填，默认1）
        "style": "Table Grid",    # str 表格样式（默认"Table Grid"=带边框，可填Word内置样式名）
        "table_align": "center",  # str 表格整体对齐（"left"/"center"/"right"，默认"center"）
        "col_widths": [2, 3, 2],  # list[float] 列宽列表（单位：厘米，长度需等于cols，默认自动分配）
        "row_heights": [1, 1],    # list[float] 行高列表（单位：厘米，长度需等于rows，默认自动分配）
        "cell_vertical_align": "center",  # str 单元格垂直对齐（"top"/"center"/"bottom"，默认"center"）
        "cell_contents": [        # list[list[dict]] 单元格内容（二维列表，每行对应一个子列表，每个元素是文本样式配置）
            [   
                {"text": "照片", "size": 14, "bold": "True", "color": "#FF0000"},
                {"text": "姓名", "size": 14, "bold": "True"},  # 行0列0
                {"text": "年龄", "size": 14, "bold": "True"},                      # 行0列1
                {"text": "性别", "size": 14, "bold": "True"}                       # 行0列2
            ],
            [   
                {"path": "照片"},
                {"text": "张三", "alignment": "center"},   # 行1列0
                {"text": "25", "alignment": "center"},    # 行1列1
                {"text": "男", "alignment": "center"}     # 行1列2
            ]
        ]
    }
}

三、图片类型（type="picture"）
{
    "type": "picture",
    "config": {
        "img_path": "路径",
        "width": 10,
        "alignment": "center",
        "spacing_before": 0.5
    }
}
四、分页符（type="page_breaks"）
{
    "type": "page_breaks",
    "config": {}
}

configs = {
    "header": {
        "enabled": True,
        "text": "公司内部文档",
        "font_name": "微软雅黑",
        "font_size": 11,
        "color": "#000000",
        "alignment": "right"
    },
    "footer": {
        "enabled": True,
        "text": "机密文件，请勿外传",
        "font_name": "宋体",
        "font_size": 9,
        "color": "#000000",
        "alignment": "right"
    },
    "page_number": {
        "enabled": True,
        "format": "第 X 页",  # X 自动替换为页码
        "font_name": "Times New Roman",
        "font_size": 9,
        "color": "#000000",
        "alignment": "center"
    },
    "items": [其他段落配置]
}
"""

class DocumentDocx:
    def __init__(self):
        self.doc = None
        self.align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        self.underline_map = {
            "False": WD_UNDERLINE.NONE,
            "single": WD_UNDERLINE.SINGLE,
            "double": WD_UNDERLINE.DOUBLE,
            "dotted": WD_UNDERLINE.DOTTED,
            "dash": WD_UNDERLINE.DASH,
        }
        self.cell_vert_align_map = {
            "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
            "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
        }

    def _hex_to_rgb(self, hex_str: str) -> RGBColor:
        hex_str = hex_str.lstrip('#')
        if len(hex_str) == 3:
            hex_str = ''.join(c * 2 for c in hex_str)
        if len(hex_str) != 6:
            raise ValueError(f"颜色格式错误：{hex_str}")
        return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

    def _write_paragraph(self, para_config: Dict[str, Any], content_list: List[Dict[str, Any]]):
        p = self.doc.add_paragraph()

        line_spacing = para_config.get("line_spacing", 1)
        if isinstance(line_spacing, (int, float)):
            if line_spacing == 1:
                p.paragraph_format.line_spacing = WD_LINE_SPACING.SINGLE
            elif line_spacing == 1.5:
                p.paragraph_format.line_spacing = WD_LINE_SPACING.ONE_POINT_FIVE
            elif line_spacing == 2:
                p.paragraph_format.line_spacing = WD_LINE_SPACING.DOUBLE
            else:
                p.paragraph_format.line_spacing = Pt(line_spacing)

        first_line_indent = para_config.get("first_line_indent", 0)
        if first_line_indent > 0:
            p.paragraph_format.first_line_indent = Cm(first_line_indent * 0.35)

        space_before = para_config.get("space_before", 0)
        if space_before > 0:
            p.paragraph_format.space_before = Cm(space_before)
        space_after = para_config.get("space_after", 0)
        if space_after > 0:
            p.paragraph_format.space_after = Cm(space_after)

        for content in content_list:
            text = content.get("text", "")
            if not text:
                continue
            color = content.get("color", "#000000")
            font_name = content.get("name", "宋体")
            font_size = content.get("size", 12)
            bold = content.get("bold", "False")
            italic = content.get("italic", "False")
            underline = content.get("underline", "False")
            alignment = content.get("alignment", "left")

            run = p.add_run(text)
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
            run.font.color.rgb = self._hex_to_rgb(color)
            run.bold = (True if bold.lower() == "true" else False) if isinstance(bold, str) else bold
            run.italic = (True if italic.lower() == "true" else False) if isinstance(italic, str) else italic
            run.font.underline = self.underline_map.get(underline.lower(), WD_UNDERLINE.NONE)

            p.alignment = self.align_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)

    # def _write_table(self, table_config: Dict[str, Any]):
    #     rows = table_config.get("rows", 1)
    #     cols = table_config.get("cols", 1)
    #     table_style = table_config.get("style", "Table Grid")
    #     table_align = table_config.get("table_align", "center")
    #     col_widths = table_config.get("col_widths", [])
    #     row_heights = table_config.get("row_heights", [])
    #     cell_vertical_align = table_config.get("cell_vertical_align", "center")
    #     cell_contents = table_config.get("cell_contents", [])
    #
    #     table = self.doc.add_table(rows=rows, cols=cols, style=table_style)
    #     table.alignment = self.align_map.get(table_align.lower(), WD_TABLE_ALIGNMENT.CENTER)
    #
    #     if col_widths and len(col_widths) == cols:
    #         for idx, width in enumerate(col_widths):
    #             table.columns[idx].width = Cm(width)
    #
    #     for row_idx in range(rows):
    #         row = table.rows[row_idx]
    #         if row_heights and len(row_heights) == rows:
    #             row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    #             row.height = Cm(row_heights[row_idx])
    #
    #         for col_idx in range(cols):
    #             cell = table.cell(row_idx, col_idx)
    #             cell.vertical_alignment = self.cell_vert_align_map.get(
    #                 cell_vertical_align.lower(), WD_CELL_VERTICAL_ALIGNMENT.CENTER
    #             )
    #             cell_text_config = {}
    #             if cell_contents and len(cell_contents) > row_idx and len(cell_contents[row_idx]) > col_idx:
    #                 cell_text_config = cell_contents[row_idx][col_idx]
    #
    #             cell.text = ""
    #             if cell_text_config:
    #                 p = cell.add_paragraph()
    #                 run = p.add_run(cell_text_config.get("text", ""))
    #                 font_name = cell_text_config.get("name", "宋体")
    #                 run.font.name = font_name
    #                 run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    #                 run.font.size = Pt(cell_text_config.get("size", 12))
    #                 run.font.color.rgb = self._hex_to_rgb(cell_text_config.get("color", "#000000"))
    #                 run.bold = (cell_text_config.get("bold", "False").lower() == "true")
    #                 run.italic = (cell_text_config.get("italic", "False").lower() == "true")
    #                 p.alignment = self.align_map.get(
    #                     cell_text_config.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER
    #                 )

    def _write_table(self, table_config: Dict[str, Any]):
        rows = table_config.get("rows", 1)
        cols = table_config.get("cols", 1)
        table_style = table_config.get("style", "Table Grid")
        table_align = table_config.get("table_align", "center")
        col_widths = table_config.get("col_widths", [])
        row_heights = table_config.get("row_heights", [])
        cell_vertical_align = table_config.get("cell_vertical_align", "center")
        cell_contents = table_config.get("cell_contents", [])

        table = self.doc.add_table(rows=rows, cols=cols, style=table_style)
        table.alignment = self.align_map.get(table_align.lower(), WD_TABLE_ALIGNMENT.CENTER)

        if col_widths and len(col_widths) == cols:
            for idx, width in enumerate(col_widths):
                table.columns[idx].width = Cm(width)

        for row_idx in range(rows):
            row = table.rows[row_idx]
            if row_heights and len(row_heights) == rows:
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                row.height = Cm(row_heights[row_idx])

            for col_idx in range(cols):
                cell = table.cell(row_idx, col_idx)
                cell.vertical_alignment = self.cell_vert_align_map.get(
                    cell_vertical_align.lower(), WD_CELL_VERTICAL_ALIGNMENT.CENTER
                )

                cell_text_config = {}
                if cell_contents and len(cell_contents) > row_idx and len(cell_contents[row_idx]) > col_idx:
                    cell_text_config = cell_contents[row_idx][col_idx]

                cell.text = ""

                # ----- 判断是文本还是图片 -----
                if "text" in cell_text_config or "path" not in cell_text_config:
                    # 文本单元格（原逻辑）
                    if cell_text_config:
                        p = cell.add_paragraph()
                        run = p.add_run(cell_text_config.get("text", ""))
                        font_name = cell_text_config.get("name", "宋体")
                        run.font.name = font_name
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                        run.font.size = Pt(cell_text_config.get("size", 12))
                        run.font.color.rgb = self._hex_to_rgb(cell_text_config.get("color", "#000000"))
                        run.bold = (cell_text_config.get("bold", "False").lower() == "true")
                        run.italic = (cell_text_config.get("italic", "False").lower() == "true")
                        p.alignment = self.align_map.get(
                            cell_text_config.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER
                        )
                else:
                    # 图片单元格
                    img_path = cell_text_config["path"]
                    p = cell.add_paragraph()
                    run = p.add_run()

                    # 计算可用尺寸（保留少量边距）
                    margin_cm = 0.2
                    avail_width = cell.width - Cm(margin_cm)  # Length 对象
                    avail_height = None
                    if row_heights and len(row_heights) > row_idx:
                        avail_height = Cm(row_heights[row_idx]) - Cm(margin_cm)

                    # 用户手动指定的尺寸优先
                    user_width = cell_text_config.get("width")
                    user_height = cell_text_config.get("height")

                    if user_width is not None:
                        run.add_picture(img_path, width=Cm(user_width))
                    elif user_height is not None:
                        run.add_picture(img_path, height=Cm(user_height))
                    else:
                        # 尝试自适应缩放（需要 Pillow）
                        try:
                            from PIL import Image as PILImage
                            with PILImage.open(img_path) as img:
                                orig_w, orig_h = img.size
                                # 计算可用尺寸（转换为厘米数值）
                                avail_w_cm = avail_width / 360000  # EMU → cm
                                # 以宽度为基准缩放
                                scale = avail_w_cm / orig_w
                                # 如果高度受限，补偿
                                if avail_height is not None:
                                    avail_h_cm = avail_height / 360000
                                    scale = min(scale, avail_h_cm / orig_h)
                                # 应用缩放（保留一些边距，可微调）
                                final_w_cm = orig_w * scale * 0.95  # 留 5% 边距
                                run.add_picture(img_path, width=Cm(final_w_cm))
                        except ImportError:
                            # 无 Pillow，使用默认宽度
                            run.add_picture(img_path, width=Cm(3))
                    # 图片段落对齐
                    p.alignment = self.align_map.get(
                        cell_text_config.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER
                    )

    def _write_picture(self, pic_config: Dict[str, Any]):
        img_path = pic_config.get("img_path", "")
        if not img_path or not os.path.exists(img_path):
            raise FileNotFoundError(f"图片路径不存在：{img_path}")

        width = pic_config.get("width")
        height = pic_config.get("height")
        alignment = pic_config.get("alignment", "center")
        spacing_before = pic_config.get("spacing_before", 0)
        spacing_after = pic_config.get("spacing_after", 0)

        p = self.doc.add_paragraph()
        p.alignment = self.align_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_before = Cm(spacing_before)
        p.paragraph_format.space_after = Cm(spacing_after)

        run = p.add_run()
        if width:
            run.add_picture(img_path, width=Cm(width))
        elif height:
            run.add_picture(img_path, height=Cm(height))
        else:
            run.add_picture(img_path, width=Cm(15))

    def _page_breaks(self, config: Dict[str, Any]):
        p = self.doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

    def _setup_header(self, config: dict):
        if not config.get("enabled", False):
            return
        section = self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.clear()
        p.alignment = self.align_map.get(config.get("alignment", "center").lower(), WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(config.get("text", ""))
        run.font.name = config.get("font_name", "宋体")
        run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
        run.font.size = Pt(config.get("font_size", 10))
        run.font.color.rgb = self._hex_to_rgb(config.get("color", "#808080"))

    def _setup_footer(self, config: dict):
        if not config.get("enabled", False):
            return
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.clear()
        p.alignment = self.align_map.get(config.get("alignment", "center").lower(), WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(config.get("text", ""))
        run.font.name = config.get("font_name", "宋体")
        run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
        run.font.size = Pt(config.get("font_size", 9))
        run.font.color.rgb = self._hex_to_rgb(config.get("color", "#808080"))

    def _setup_page_number(self, config: dict):
        if not config.get("enabled", False):
            return
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs[0].text.strip() != "":
            p = footer.add_paragraph()
        else:
            p = footer.paragraphs[0]
            p.clear()
        p.alignment = self.align_map.get(config.get("alignment", "center").lower(), WD_ALIGN_PARAGRAPH.CENTER)

        template = config.get("format", "第 X 页")
        parts = template.split("X")
        for i, part in enumerate(parts):
            if part:
                run = p.add_run(part)
                run.font.name = config.get("font_name", "宋体")
                run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
                run.font.size = Pt(config.get("font_size", 9))
                run.font.color.rgb = self._hex_to_rgb(config.get("color", "#000000"))
            if i < len(parts) - 1:
                run_page = p.add_run()
                run_page.font.name = config.get("font_name", "宋体")
                run_page._element.rPr.rFonts.set(qn('w:eastAsia'), run_page.font.name)
                run_page.font.size = Pt(config.get("font_size", 9))
                run_page.font.color.rgb = self._hex_to_rgb(config.get("color", "#000000"))
                fldChar_begin = OxmlElement('w:fldChar')
                fldChar_begin.set(qn('w:fldCharType'), 'begin')
                run_page._element.append(fldChar_begin)
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = ' PAGE '
                run_page._element.append(instrText)
                fldChar_end = OxmlElement('w:fldChar')
                fldChar_end.set(qn('w:fldCharType'), 'end')
                run_page._element.append(fldChar_end)

    def write_doc(self, configs: dict, save_path: str):
        self.doc = Document()
        header_config = configs.get("header", {})
        footer_config = configs.get("footer", {})
        page_number_config = configs.get("page_number", {})
        self._setup_header(header_config)
        self._setup_footer(footer_config)
        self._setup_page_number(page_number_config)

        items = configs.get("items", [])
        for content_data in items:
            content_type = content_data.get("type", "")
            config = content_data.get("config", {})
            if content_type == "paragraph":
                self._write_paragraph(config, content_data.get("content", []))
            elif content_type == "table":
                self._write_table(config)
            elif content_type == "picture":
                self._write_picture(config)
            elif content_type == "page_breaks":
                self._page_breaks(config)

        self.doc.save(save_path)
        print(f"文档已生成：{save_path}")

if __name__ == '__main__':
    configs = {
        "header": {
            "enabled": True,
            "text": "公司内部文档",
            "font_name": "微软雅黑",
            "font_size": 11,
            "color": "#000000",
            "alignment": "right"
        },
        "footer": {
            "enabled": True,
            "text": "机密文件，请勿外传",
            "font_name": "宋体",
            "font_size": 9,
            "color": "#000000",
            "alignment": "right"
        },
        "page_number": {
            "enabled": True,
            "format": "第 X 页",  # X 自动替换为页码
            "font_name": "Times New Roman",
            "font_size": 9,
            "color": "#000000",
            "alignment": "center"
        },
        "items": [{
            "type": "table",
            "config": {
                "rows": 2,
                "cols": 3,
                "cell_vertical_align": "center",
                "cell_contents": [
                    [
                        {"text": "照片", "size": 14, "bold": "True", "color": "#FF0000"},
                        {"text": "姓名", "size": 14, "bold": "True"},
                        {"text": "年龄", "size": 14, "bold": "True"},
                        {"text": "性别", "size": 14, "bold": "True"}
                    ],
                    [
                        {"path": "0.png"},
                        {"text": "张三", "alignment": "center"},
                        {"text": "25", "alignment": "center"},
                        {"text": "男", "alignment": "center"}
                    ]
                ]
            }
        }]
    }

    d = DocumentDocx()
    d.write_doc(configs, 'n.docx')